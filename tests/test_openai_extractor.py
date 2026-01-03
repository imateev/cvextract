"""Tests for OpenAI CV extractor implementation."""

import json
import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock
from cvextract.extractors import OpenAICVExtractor, CVExtractor


class TestOpenAICVExtractor:
    """Tests for OpenAICVExtractor implementation."""

    def test_openai_extractor_is_cv_extractor(self):
        """OpenAICVExtractor is a CVExtractor."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            assert isinstance(extractor, CVExtractor)

    def test_openai_extractor_default_model(self):
        """OpenAICVExtractor uses default model."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            assert extractor.model == 'gpt-4o'

    def test_openai_extractor_custom_model(self):
        """OpenAICVExtractor accepts custom model."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor(model='gpt-4-turbo')
            assert extractor.model == 'gpt-4-turbo'

    def test_extract_raises_file_not_found(self):
        """extract() raises FileNotFoundError for missing file."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            with pytest.raises(FileNotFoundError):
                extractor.extract(Path('/nonexistent/file.pdf'))

    def test_extract_raises_value_error_for_directory(self, tmp_path):
        """extract() raises ValueError for directory."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            test_dir = tmp_path / "test_dir"
            test_dir.mkdir()
            with pytest.raises(ValueError, match="must be a file"):
                extractor.extract(test_dir)

    def test_extract_accepts_any_file_type(self, tmp_path):
        """extract() accepts any file type without validation."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            test_file = tmp_path / "test.xyz"
            test_file.write_text("test content")
            
            # Mock OpenAI response
            mock_response = {
                "identity": {
                    "title": "Test",
                    "full_name": "Test User",
                    "first_name": "Test",
                    "last_name": "User"
                },
                "sidebar": {},
                "overview": "Test",
                "experiences": []
            }
            
            with patch.object(extractor, '_extract_with_openai', return_value=json.dumps(mock_response)):
                result = extractor.extract(test_file)
                assert result['identity']['full_name'] == 'Test User'

    def test_extract_supports_txt_files(self, tmp_path):
        """extract() supports .txt files via OpenAI API."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Create test file
            test_file = tmp_path / "test.txt"
            test_file.write_text("Test CV content\nJohn Doe\nSoftware Engineer")
            
            # Mock OpenAI response
            mock_response = {
                "identity": {
                    "title": "Software Engineer",
                    "full_name": "John Doe",
                    "first_name": "John",
                    "last_name": "Doe"
                },
                "sidebar": {
                    "languages": ["Python"],
                    "tools": [],
                    "certifications": [],
                    "industries": [],
                    "spoken_languages": ["English"],
                    "academic_background": []
                },
                "overview": "Experienced software engineer",
                "experiences": []
            }
            
            with patch.object(extractor, '_extract_with_openai', return_value=json.dumps(mock_response)):
                result = extractor.extract(test_file)
                assert result['identity']['full_name'] == 'John Doe'

    def test_extract_supports_pdf_files(self, tmp_path):
        """extract() supports .pdf files via OpenAI API."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Create test PDF file
            test_file = tmp_path / "test.pdf"
            test_file.write_bytes(b'%PDF-1.4\ntest content')
            
            # Mock OpenAI response
            mock_response = {
                "identity": {
                    "title": "Engineer",
                    "full_name": "Jane PDF",
                    "first_name": "Jane",
                    "last_name": "PDF"
                },
                "sidebar": {},
                "overview": "PDF CV",
                "experiences": []
            }
            
            with patch.object(extractor, '_extract_with_openai', return_value=json.dumps(mock_response)):
                result = extractor.extract(test_file)
                assert result['identity']['full_name'] == 'Jane PDF'

    def test_extract_supports_docx_files(self, tmp_path):
        """extract() supports .docx files via OpenAI API."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Create a real DOCX file using python-docx
            from docx import Document
            doc = Document()
            doc.add_paragraph("Jane Smith")
            doc.add_paragraph("Developer")
            doc.add_paragraph("Experienced developer with Python skills")
            
            test_file = tmp_path / "test.docx"
            doc.save(str(test_file))
            
            # Mock OpenAI response
            mock_response = {
                "identity": {
                    "title": "Developer",
                    "full_name": "Jane Smith",
                    "first_name": "Jane",
                    "last_name": "Smith"
                },
                "sidebar": {},
                "overview": "Developer overview",
                "experiences": []
            }
            
            with patch.object(extractor, '_extract_with_openai', return_value=json.dumps(mock_response)):
                result = extractor.extract(test_file)
                assert result['identity']['full_name'] == 'Jane Smith'

    def test_extract_supports_pptx_files(self, tmp_path):
        """extract() supports .pptx files via OpenAI API."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Create test PPTX file
            test_file = tmp_path / "test.pptx"
            test_file.write_bytes(b'PK\x03\x04test pptx content')
            
            # Mock OpenAI response
            mock_response = {
                "identity": {
                    "title": "Manager",
                    "full_name": "PPTX User",
                    "first_name": "PPTX",
                    "last_name": "User"
                },
                "sidebar": {},
                "overview": "PPTX CV",
                "experiences": []
            }
            
            with patch.object(extractor, '_extract_with_openai', return_value=json.dumps(mock_response)):
                result = extractor.extract(test_file)
                assert result['identity']['full_name'] == 'PPTX User'

    def test_parse_response_handles_markdown_code_blocks(self, tmp_path):
        """_parse_and_validate_response() handles markdown code blocks."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Create a response with markdown code block
            response = """```json
{
    "identity": {"title": "Test", "full_name": "Test", "first_name": "T", "last_name": "est"},
    "sidebar": {},
    "overview": "Test",
    "experiences": []
}
```"""
            
            # Load schema
            schema_path = Path(__file__).parent.parent / "cvextract" / "contracts" / "cv_schema.json"
            with open(schema_path, 'r') as f:
                schema = json.load(f)
            
            result = extractor._parse_and_validate_response(response, schema)
            assert result['identity']['full_name'] == 'Test'

    def test_parse_response_validates_required_fields(self, tmp_path):
        """_parse_and_validate_response() validates against schema using CVSchemaVerifier."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Incomplete response missing required fields
            response = '{"identity": {"title": "Test", "full_name": "Test", "first_name": "T", "last_name": "est"}}'
            
            # Load schema
            schema_path = Path(__file__).parent.parent / "cvextract" / "contracts" / "cv_schema.json"
            with open(schema_path, 'r') as f:
                schema = json.load(f)
            
            # Should raise ValueError for missing required fields (sidebar, overview, experiences)
            with pytest.raises(ValueError, match="Extracted data failed schema validation"):
                extractor._parse_and_validate_response(response, schema)

    def test_extract_full_integration_with_mock(self, tmp_path):
        """Full integration test with mocked OpenAI API."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            # Create test file
            test_file = tmp_path / "cv.txt"
            test_file.write_text("""
John Doe
Senior Software Engineer

Skills: Python, Java, JavaScript
Languages: English, Spanish

Overview:
Experienced software engineer with 10 years in the industry.

Experience:
2020-Present | Senior Engineer at TechCorp
- Led development of microservices
- Managed team of 5 engineers

2015-2020 | Software Engineer at StartupCo
- Developed web applications
- Improved system performance
""")
            
            # Mock the OpenAI client
            mock_client = MagicMock()
            mock_response = MagicMock()
            mock_response.choices = [MagicMock()]
            mock_response.choices[0].message.content = json.dumps({
                "identity": {
                    "title": "Senior Software Engineer",
                    "full_name": "John Doe",
                    "first_name": "John",
                    "last_name": "Doe"
                },
                "sidebar": {
                    "languages": ["Python", "Java", "JavaScript"],
                    "tools": [],
                    "certifications": [],
                    "industries": ["Technology"],
                    "spoken_languages": ["English", "Spanish"],
                    "academic_background": []
                },
                "overview": "Experienced software engineer with 10 years in the industry.",
                "experiences": [
                    {
                        "heading": "2020-Present | Senior Engineer at TechCorp",
                        "description": "Led development of microservices",
                        "bullets": ["Led development of microservices", "Managed team of 5 engineers"],
                        "environment": ["Microservices"]
                    },
                    {
                        "heading": "2015-2020 | Software Engineer at StartupCo",
                        "description": "Developed web applications",
                        "bullets": ["Developed web applications", "Improved system performance"],
                        "environment": None
                    }
                ]
            })
            mock_client.chat.completions.create.return_value = mock_response
            
            extractor = OpenAICVExtractor()
            extractor.client = mock_client
            
            # Extract - OpenAI API receives the file path, not content
            result = extractor.extract(test_file)
            
            # Verify result
            assert result['identity']['full_name'] == 'John Doe'
            assert result['identity']['title'] == 'Senior Software Engineer'
            assert 'Python' in result['sidebar']['languages']
            assert 'English' in result['sidebar']['spoken_languages']
            assert len(result['experiences']) == 2
            assert result['experiences'][0]['heading'].startswith('2020-Present')

    def test_extract_handles_openai_api_error(self, tmp_path):
        """extract() handles OpenAI API errors gracefully."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Create test file
            test_file = tmp_path / "test.txt"
            test_file.write_text("Test CV content")
            
            # Mock _extract_with_openai to raise an exception
            with patch.object(extractor, '_extract_with_openai', side_effect=Exception("API Error")):
                with pytest.raises(Exception, match="API Error"):
                    extractor.extract(test_file)

    def test_parse_response_handles_plain_markdown_blocks(self, tmp_path):
        """_parse_and_validate_response() handles plain markdown blocks (without json label)."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Create a response with plain markdown code block (no json label)
            response = """```
{
    "identity": {"title": "Test", "full_name": "Test User", "first_name": "Test", "last_name": "User"},
    "sidebar": {},
    "overview": "Test overview",
    "experiences": []
}
```"""
            
            # Load schema
            schema_path = Path(__file__).parent.parent / "cvextract" / "contracts" / "cv_schema.json"
            with open(schema_path, 'r') as f:
                schema = json.load(f)
            
            result = extractor._parse_and_validate_response(response, schema)
            assert result['identity']['full_name'] == 'Test User'

    def test_parse_response_handles_invalid_json(self, tmp_path):
        """_parse_and_validate_response() raises error for invalid JSON."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Invalid JSON response
            response = '{"identity": invalid json here}'
            
            # Load schema
            schema_path = Path(__file__).parent.parent / "cvextract" / "contracts" / "cv_schema.json"
            with open(schema_path, 'r') as f:
                schema = json.load(f)
            
            with pytest.raises(ValueError, match="OpenAI returned invalid JSON"):
                extractor._parse_and_validate_response(response, schema)

    def test_parse_response_validates_identity_type(self, tmp_path):
        """_parse_and_validate_response() validates identity structure via CVSchemaVerifier."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Response with invalid identity (not a dict)
            response = '{"identity": "not a dict", "sidebar": {}, "overview": "", "experiences": []}'
            
            # Load schema
            schema_path = Path(__file__).parent.parent / "cvextract" / "contracts" / "cv_schema.json"
            with open(schema_path, 'r') as f:
                schema = json.load(f)
            
            # Should raise ValueError for invalid identity type
            with pytest.raises(ValueError, match="Extracted data failed schema validation"):
                extractor._parse_and_validate_response(response, schema)

    def test_extract_raises_runtime_error_when_system_prompt_fails_to_load(self, tmp_path):
        """_extract_with_openai() raises RuntimeError when system prompt fails to load."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Create test file
            test_file = tmp_path / "test.txt"
            test_file.write_text("Test CV content")
            
            # Mock load_prompt to return None for system prompt
            with patch('cvextract.extractors.openai_extractor.load_prompt', return_value=None):
                with pytest.raises(RuntimeError, match="Failed to load CV extraction system prompt"):
                    extractor.extract(test_file)

    def test_extract_raises_runtime_error_when_user_prompt_fails_to_format(self, tmp_path):
        """_extract_with_openai() raises RuntimeError when user prompt fails to format."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Create test file
            test_file = tmp_path / "test.txt"
            test_file.write_text("Test CV content")
            
            # Mock load_prompt to succeed but format_prompt to fail
            with patch('cvextract.extractors.openai_extractor.load_prompt', return_value="System prompt"):
                with patch('cvextract.extractors.openai_extractor.format_prompt', return_value=None):
                    with pytest.raises(RuntimeError, match="Failed to format CV extraction user prompt"):
                        extractor.extract(test_file)

    def test_extract_raises_exception_when_openai_api_fails(self, tmp_path):
        """_extract_with_openai() raises Exception when OpenAI API call fails."""
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'test-key'}):
            extractor = OpenAICVExtractor()
            
            # Create test file
            test_file = tmp_path / "test.txt"
            test_file.write_text("Test CV content")
            
            # Mock the OpenAI client to raise an exception
            mock_client = MagicMock()
            mock_client.chat.completions.create.side_effect = Exception("API connection failed")
            extractor.client = mock_client
            
            with pytest.raises(Exception, match="Failed to send document to OpenAI: API connection failed"):
                extractor.extract(test_file)
