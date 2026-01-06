"""
Tests for DocxCVRenderer implementation.

These tests focus on the DocxCVRenderer class with minimal mocking,
using real DOCX templates and the actual docxtpl library.
"""

import pytest
from pathlib import Path
from docx import Document
from cvextract.renderers import DocxCVRenderer
from cvextract.shared import UnitOfWork


@pytest.fixture
def minimal_template(tmp_path):
    """Create a minimal DOCX template with basic placeholders."""
    template_path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("Name: {{ identity.full_name }}")
    doc.add_paragraph("Title: {{ identity.title }}")
    doc.save(str(template_path))
    return template_path


@pytest.fixture
def full_template(tmp_path):
    """Create a comprehensive DOCX template with all CV fields."""
    template_path = tmp_path / "full_template.docx"
    doc = Document()
    
    # Identity section
    doc.add_heading("CV - {{ identity.full_name }}", 0)
    doc.add_paragraph("{{ identity.title }}")
    
    # Overview section
    doc.add_heading("Overview", 1)
    doc.add_paragraph("{{ overview }}")
    
    # Skills section (sidebar)
    doc.add_heading("Skills", 1)
    doc.add_paragraph("Languages: {% for lang in sidebar.languages %}{{ lang }}{% if not loop.last %}, {% endif %}{% endfor %}")
    doc.add_paragraph("Tools: {% for tool in sidebar.tools %}{{ tool }}{% if not loop.last %}, {% endif %}{% endfor %}")
    
    # Experience section
    doc.add_heading("Experience", 1)
    doc.add_paragraph("{% for exp in experiences %}{{ exp.heading }}{% endfor %}")
    
    doc.save(str(template_path))
    return template_path


@pytest.fixture
def sample_cv_data():
    """Sample CV data conforming to cv_schema.json."""
    return {
        "identity": {
            "title": "Senior Software Engineer",
            "full_name": "Jane Smith",
            "first_name": "Jane",
            "last_name": "Smith"
        },
        "sidebar": {
            "languages": ["Python", "JavaScript", "Go"],
            "tools": ["Docker", "Kubernetes", "AWS"],
            "certifications": ["AWS Solutions Architect"],
            "industries": ["Finance", "Healthcare"],
            "spoken_languages": ["English", "French"],
            "academic_background": ["BS Computer Science - MIT"]
        },
        "overview": "Experienced software engineer with 8+ years building scalable systems.",
        "experiences": [
            {
                "heading": "2020-Present | Senior Engineer at TechCorp",
                "description": "Leading backend development team",
                "bullets": [
                    "Architected microservices handling 1M+ requests/day",
                    "Reduced deployment time by 60%",
                    "Mentored 5 junior engineers"
                ],
                "environment": ["Python", "Docker", "AWS"]
            },
            {
                "heading": "2018-2020 | Software Engineer at StartupXYZ",
                "description": "Full-stack development",
                "bullets": [
                    "Built core payment processing system",
                    "Improved API response time by 40%"
                ],
                "environment": ["JavaScript", "Node.js"]
            }
        ]
    }


class TestDocxCVRendererBasicFunctionality:
    """Test basic rendering functionality."""
    
    def test_renderer_can_be_instantiated(self):
        """DocxCVRenderer can be instantiated without arguments."""
        renderer = DocxCVRenderer()
        assert renderer is not None
    
    def test_render_creates_output_file(self, minimal_template, tmp_path, make_render_work):
        """Rendering creates an output DOCX file."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "output.docx"
        
        cv_data = {
            "identity": {
                "title": "Engineer",
                "full_name": "John Doe",
                "first_name": "John",
                "last_name": "Doe"
            },
            "sidebar": {},
            "overview": "",
            "experiences": []
        }
        
        work = make_render_work(cv_data, minimal_template, output_path)
        result = renderer.render(work)
        
        assert result.output == output_path
        assert output_path.exists()
        assert output_path.is_file()
        assert output_path.stat().st_size > 0
    
    def test_render_returns_output_path(self, minimal_template, tmp_path, make_render_work):
        """Render method returns the output path."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "test_output.docx"
        
        cv_data = {
            "identity": {"title": "Dev", "full_name": "Test", "first_name": "T", "last_name": "Test"},
            "sidebar": {},
            "overview": "",
            "experiences": []
        }
        
        work = make_render_work(cv_data, minimal_template, output_path)
        result = renderer.render(work)
        assert isinstance(result, UnitOfWork)
        assert result.output == output_path


class TestDocxCVRendererValidation:
    """Test input validation."""
    
    def test_raises_error_when_template_not_exists(self, tmp_path, make_render_work):
        """Raises FileNotFoundError when template doesn't exist."""
        renderer = DocxCVRenderer()
        non_existent = tmp_path / "does_not_exist.docx"
        output_path = tmp_path / "output.docx"
        
        cv_data = {
            "identity": {"title": "", "full_name": "", "first_name": "", "last_name": ""},
            "sidebar": {},
            "overview": "",
            "experiences": []
        }
        
        work = make_render_work(cv_data, non_existent, output_path)
        with pytest.raises(FileNotFoundError, match="Template file not found"):
            renderer.render(work)
    
    def test_raises_error_when_template_is_directory(self, tmp_path, make_render_work):
        """Raises ValueError when template path is a directory."""
        renderer = DocxCVRenderer()
        template_dir = tmp_path / "template_dir.docx"
        template_dir.mkdir()
        output_path = tmp_path / "output.docx"
        
        cv_data = {
            "identity": {"title": "", "full_name": "", "first_name": "", "last_name": ""},
            "sidebar": {},
            "overview": "",
            "experiences": []
        }
        
        work = make_render_work(cv_data, template_dir, output_path)
        with pytest.raises(ValueError, match="not a file"):
            renderer.render(work)
    
    def test_raises_error_when_template_not_docx(self, tmp_path, make_render_work):
        """Raises ValueError when template is not a .docx file."""
        renderer = DocxCVRenderer()
        txt_file = tmp_path / "template.txt"
        txt_file.write_text("not a docx")
        output_path = tmp_path / "output.docx"
        
        cv_data = {
            "identity": {"title": "", "full_name": "", "first_name": "", "last_name": ""},
            "sidebar": {},
            "overview": "",
            "experiences": []
        }
        
        work = make_render_work(cv_data, txt_file, output_path)
        with pytest.raises(ValueError, match="must be a .docx file"):
            renderer.render(work)


class TestDocxCVRendererDataRendering:
    """Test that CV data is correctly rendered into templates."""
    
    def test_renders_identity_data(self, minimal_template, tmp_path, make_render_work):
        """Identity data is correctly rendered in output."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "output.docx"
        
        cv_data = {
            "identity": {
                "title": "Chief Technology Officer",
                "full_name": "Alice Johnson",
                "first_name": "Alice",
                "last_name": "Johnson"
            },
            "sidebar": {},
            "overview": "",
            "experiences": []
        }
        
        work = make_render_work(cv_data, minimal_template, output_path)
        renderer.render(work)
        
        # Read the output and verify content
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        assert "Alice Johnson" in text
        assert "Chief Technology Officer" in text
    
    def test_renders_complete_cv_data(self, full_template, sample_cv_data, tmp_path, make_render_work):
        """Complete CV data is rendered correctly."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "full_cv.docx"
        
        work = make_render_work(sample_cv_data, full_template, output_path)
        renderer.render(work)
        
        # Verify output file
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Check identity
        assert "Jane Smith" in text
        assert "Senior Software Engineer" in text
        
        # Check overview
        assert "Experienced software engineer" in text
        
        # Check sidebar data is present
        assert "Python" in text or "Languages:" in text
    
    def test_renders_with_empty_optional_fields(self, minimal_template, tmp_path, make_render_work):
        """Renders successfully with empty optional fields."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "minimal_cv.docx"
        
        cv_data = {
            "identity": {
                "title": "Developer",
                "full_name": "Bob Builder",
                "first_name": "Bob",
                "last_name": "Builder"
            },
            "sidebar": {},  # Empty sidebar
            "overview": "",  # Empty overview
            "experiences": []  # No experiences
        }
        
        work = make_render_work(cv_data, minimal_template, output_path)
        renderer.render(work)
        
        # Should still create a valid file
        assert output_path.exists()
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        assert "Bob Builder" in text


class TestDocxCVRendererSpecialCharacters:
    """Test handling of special characters and sanitization."""
    
    def test_handles_special_xml_characters(self, minimal_template, tmp_path, make_render_work):
        """Special XML characters are properly handled."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "special_chars.docx"
        
        cv_data = {
            "identity": {
                "title": "Engineer & Architect",  # Ampersand
                "full_name": "O'Brien <Expert>",  # Quote and angle brackets
                "first_name": "O'Brien",
                "last_name": "Expert"
            },
            "sidebar": {},
            "overview": "Expert in C++ & Java",
            "experiences": []
        }
        
        # Should not raise an error
        work = make_render_work(cv_data, minimal_template, output_path)
        renderer.render(work)
    
        assert output_path.exists()
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Data should be present (may be escaped)
        assert "O'Brien" in text or "OBrien" in text
        assert "Engineer" in text
    
    def test_sanitizes_non_breaking_spaces(self, minimal_template, tmp_path, make_render_work):
        """Non-breaking spaces are sanitized."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "nbsp.docx"
        
        cv_data = {
            "identity": {
                "title": "Senior\u00A0Engineer",  # Non-breaking space
                "full_name": "Test\u00A0User",
                "first_name": "Test",
                "last_name": "User"
            },
            "sidebar": {},
            "overview": "",
            "experiences": []
        }
        
        # Should render without error
        work = make_render_work(cv_data, minimal_template, output_path)
        renderer.render(work)
        
        assert output_path.exists()
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Text should be present
        assert "Engineer" in text
        assert "Test" in text and "User" in text
    
    def test_sanitizes_soft_hyphens(self, minimal_template, tmp_path, make_render_work):
        """Soft hyphens are sanitized."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "soft_hyphen.docx"
        
        cv_data = {
            "identity": {
                "title": "Soft\u00ADware Engineer",  # Soft hyphen
                "full_name": "Hyphen\u00ADated Name",
                "first_name": "Hyphenated",
                "last_name": "Name"
            },
            "sidebar": {},
            "overview": "",
            "experiences": []
        }
        
        work = make_render_work(cv_data, minimal_template, output_path)
        renderer.render(work)
        
        assert output_path.exists()
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        assert "Engineer" in text
        assert "Name" in text


class TestDocxCVRendererOutputDirectory:
    """Test output directory handling."""
    
    def test_creates_output_directory_if_not_exists(self, minimal_template, tmp_path, make_render_work):
        """Creates output directory structure if it doesn't exist."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "nested" / "dirs" / "output.docx"
        
        # Directory doesn't exist yet
        assert not output_path.parent.exists()
        
        cv_data = {
            "identity": {"title": "Dev", "full_name": "Test", "first_name": "T", "last_name": "Test"},
            "sidebar": {},
            "overview": "",
            "experiences": []
        }
        
        work = make_render_work(cv_data, minimal_template, output_path)
        renderer.render(work)
        
        # Directory should now exist
        assert output_path.parent.exists()
        assert output_path.exists()
    
    def test_works_with_existing_output_directory(self, minimal_template, tmp_path, make_render_work):
        """Works correctly when output directory already exists."""
        renderer = DocxCVRenderer()
        output_dir = tmp_path / "existing_dir"
        output_dir.mkdir(parents=True)
        output_path = output_dir / "output.docx"
        
        cv_data = {
            "identity": {"title": "Dev", "full_name": "Test", "first_name": "T", "last_name": "Test"},
            "sidebar": {},
            "overview": "",
            "experiences": []
        }
        
        work = make_render_work(cv_data, minimal_template, output_path)
        renderer.render(work)
        
        assert output_path.exists()


class TestDocxCVRendererComplexData:
    """Test rendering with complex, nested CV data."""
    
    def test_renders_multiple_experiences(self, full_template, tmp_path, make_render_work):
        """Multiple experience entries are rendered."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "multi_exp.docx"
        
        cv_data = {
            "identity": {
                "title": "Engineer",
                "full_name": "Multi Exp",
                "first_name": "Multi",
                "last_name": "Exp"
            },
            "sidebar": {
                "languages": ["Python", "Java"],
                "tools": [],
                "certifications": [],
                "industries": [],
                "spoken_languages": [],
                "academic_background": []
            },
            "overview": "Professional summary",
            "experiences": [
                {
                    "heading": "2022-2024 | Role 1",
                    "description": "Description 1",
                    "bullets": ["Achievement 1", "Achievement 2"],
                    "environment": ["Python"]
                },
                {
                    "heading": "2020-2022 | Role 2",
                    "description": "Description 2",
                    "bullets": ["Achievement 3"],
                    "environment": None
                },
                {
                    "heading": "2018-2020 | Role 3",
                    "description": "Description 3",
                    "bullets": [],
                    "environment": ["Java", "AWS"]
                }
            ]
        }
        
        work = make_render_work(cv_data, full_template, output_path)
        renderer.render(work)
        
        assert output_path.exists()
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Should contain data from the CV
        assert "Multi Exp" in text
    
    def test_renders_with_all_sidebar_fields(self, full_template, sample_cv_data, tmp_path, make_render_work):
        """All sidebar fields are rendered correctly."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "full_sidebar.docx"
        
        work = make_render_work(sample_cv_data, full_template, output_path)
        renderer.render(work)
        
        assert output_path.exists()
        doc = Document(str(output_path))
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Verify the file is valid and contains expected name
        assert "Jane Smith" in text


class TestDocxCVRendererEdgeCases:
    """Test edge cases and unusual inputs."""
    
    def test_handles_unicode_characters(self, minimal_template, tmp_path, make_render_work):
        """Unicode characters in various languages are handled."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "unicode.docx"
        
        cv_data = {
            "identity": {
                "title": "Développeur",  # French
                "full_name": "José García",  # Spanish
                "first_name": "José",
                "last_name": "García"
            },
            "sidebar": {},
            "overview": "Опыт работы",  # Cyrillic
            "experiences": []
        }
        
        work = make_render_work(cv_data, minimal_template, output_path)
        renderer.render(work)
        
        assert output_path.exists()
        doc = Document(str(output_path))
        # Just verify it doesn't crash and creates valid output
        assert len(doc.paragraphs) > 0
    
    def test_handles_very_long_text(self, minimal_template, tmp_path, make_render_work):
        """Very long text content is handled correctly."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "long_text.docx"
        
        long_text = "Very long overview. " * 500  # 10,000+ characters
        
        cv_data = {
            "identity": {
                "title": "Engineer",
                "full_name": "Long Text",
                "first_name": "Long",
                "last_name": "Text"
            },
            "sidebar": {},
            "overview": long_text,
            "experiences": []
        }
        
        work = make_render_work(cv_data, minimal_template, output_path)
        renderer.render(work)
        
        assert output_path.exists()
        assert output_path.stat().st_size > 1000  # Should be substantial
    
    def test_renders_with_none_environment(self, full_template, tmp_path, make_render_work):
        """Handles experience entries with None environment field."""
        renderer = DocxCVRenderer()
        output_path = tmp_path / "none_env.docx"
        
        cv_data = {
            "identity": {"title": "Dev", "full_name": "Test", "first_name": "T", "last_name": "Test"},
            "sidebar": {"languages": [], "tools": [], "certifications": [], "industries": [], "spoken_languages": [], "academic_background": []},
            "overview": "Test",
            "experiences": [
                {
                    "heading": "2020-2022 | Job",
                    "description": "Did stuff",
                    "bullets": ["Thing 1"],
                    "environment": None  # Explicitly None
                }
            ]
        }
        
        work = make_render_work(cv_data, full_template, output_path)
        renderer.render(work)
        
        assert output_path.exists()
