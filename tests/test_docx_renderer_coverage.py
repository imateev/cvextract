"""
Additional tests for DocxCVRenderer to achieve 91% coverage.

These tests cover error cases and validation paths that were missing.
"""

from pathlib import Path

import pytest

from cvextract.cli_config import RenderStage, UserConfig
from cvextract.renderers import DocxCVRenderer
from cvextract.shared import StepName, UnitOfWork


class TestDocxCVRendererErrorCases:
    """Test error cases and validation for DocxCVRenderer."""

    def test_render_raises_error_when_config_render_missing(self, tmp_path):
        """Raises ValueError when work.config.render is None."""
        renderer = DocxCVRenderer()

        # Create work with no render config
        config = UserConfig(target_dir=tmp_path)
        work = UnitOfWork(config=config)

        with pytest.raises(ValueError, match="Render configuration is missing"):
            renderer.render(work)

    def test_render_raises_error_when_output_is_none(self, tmp_path, make_render_work):
        """Raises ValueError when the render output path is None."""
        renderer = DocxCVRenderer()

        template = tmp_path / "template.docx"
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test")
        doc.save(str(template))

        cv_data = {
            "identity": {
                "title": "Dev",
                "full_name": "Test",
                "first_name": "T",
                "last_name": "Test",
            },
            "sidebar": {},
            "overview": "",
            "experiences": [],
        }

        work = make_render_work(cv_data, template, tmp_path / "output.docx")
        # Set output to None to trigger the error
        status = work.ensure_step_status(StepName.Render)
        status.output = None

        with pytest.raises(ValueError, match="Render output path is not set"):
            renderer.render(work)

    def test_render_raises_error_when_input_is_none(self, tmp_path, make_render_work):
        """Raises ValueError when the render input path is None."""
        renderer = DocxCVRenderer()

        template = tmp_path / "template.docx"
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test")
        doc.save(str(template))

        cv_data = {
            "identity": {
                "title": "Dev",
                "full_name": "Test",
                "first_name": "T",
                "last_name": "Test",
            },
            "sidebar": {},
            "overview": "",
            "experiences": [],
        }

        work = make_render_work(cv_data, template, tmp_path / "output.docx")
        # Set input to None to trigger the error
        status = work.ensure_step_status(StepName.Render)
        status.input = None

        with pytest.raises(ValueError, match="Render input path is not set"):
            renderer.render(work)
