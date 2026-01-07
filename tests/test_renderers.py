"""Tests for CV renderer interfaces and implementations."""

import json
from pathlib import Path

import pytest

from cvextract.renderers import CVRenderer, DocxCVRenderer


@pytest.fixture
def real_docx_template(tmp_path):
    """Create a real DOCX template file for integration testing."""
    from docx import Document

    template_path = tmp_path / "template.docx"
    doc = Document()

    # Add some template placeholders that match our CV schema
    doc.add_paragraph("{{ identity.full_name }}")
    doc.add_paragraph("{{ identity.title }}")
    doc.add_paragraph("{{ overview }}")

    doc.save(str(template_path))
    return template_path


class TestCVRendererInterface:
    """Tests for the CVRenderer abstract interface."""

    def test_cv_renderer_is_abstract(self):
        """CVRenderer cannot be instantiated directly."""
        with pytest.raises(TypeError):
            CVRenderer()

    def test_cv_renderer_requires_render_method(self):
        """Subclasses must implement the render method."""

        class IncompleteCVRenderer(CVRenderer):
            pass

        with pytest.raises(TypeError):
            IncompleteCVRenderer()


class TestDocxCVRenderer:
    """Tests for DocxCVRenderer implementation."""

    def test_docx_renderer_instantiation(self):
        """DocxCVRenderer can be instantiated."""
        renderer = DocxCVRenderer()
        assert isinstance(renderer, CVRenderer)
        assert isinstance(renderer, DocxCVRenderer)

    def test_render_raises_file_not_found_for_missing_template(
        self, tmp_path, make_render_work
    ):
        """render() raises FileNotFoundError for non-existent template."""
        renderer = DocxCVRenderer()
        cv_data = {
            "identity": {
                "title": "Engineer",
                "full_name": "Test User",
                "first_name": "Test",
                "last_name": "User",
            },
            "sidebar": {},
            "overview": "",
            "experiences": [],
        }
        work = make_render_work(
            cv_data, Path("/nonexistent/template.docx"), tmp_path / "out.docx"
        )
        with pytest.raises(FileNotFoundError):
            renderer.render(work)

    def test_render_raises_value_error_for_non_docx_template(
        self, tmp_path, make_render_work
    ):
        """render() raises ValueError for non-.docx template files."""
        renderer = DocxCVRenderer()
        txt_file = tmp_path / "template.txt"
        txt_file.write_text("test")
        cv_data = {
            "identity": {
                "title": "Engineer",
                "full_name": "Test User",
                "first_name": "Test",
                "last_name": "User",
            },
            "sidebar": {},
            "overview": "",
            "experiences": [],
        }
        work = make_render_work(cv_data, txt_file, tmp_path / "out.docx")
        with pytest.raises(ValueError, match="must be a .docx file"):
            renderer.render(work)

    def test_docx_renderer_implements_cv_renderer(self):
        """DocxCVRenderer properly implements CVRenderer interface."""
        renderer = DocxCVRenderer()
        assert hasattr(renderer, "render")
        assert callable(renderer.render)


class TestRendererPluggability:
    """Tests for renderer pluggability and interchangeability."""

    def test_custom_renderer_can_be_created(self, tmp_path, make_render_work):
        """Custom renderers can be created by implementing CVRenderer."""

        class MockCVRenderer(CVRenderer):
            def __init__(self):
                self.last_render = None

            def render(self, work):
                with work.input.open("r", encoding="utf-8") as f:
                    cv_data = json.load(f)
                self.last_render = {
                    "cv_data": cv_data,
                    "template": work.config.render.template,
                    "output": work.output,
                }
                # Create a mock output file
                work.output.parent.mkdir(parents=True, exist_ok=True)
                work.output.write_text("Mock rendered content")
                return work

        renderer = MockCVRenderer()
        cv_data = {
            "identity": {
                "title": "Mock Title",
                "full_name": "Mock Name",
                "first_name": "Mock",
                "last_name": "Name",
            },
            "sidebar": {},
            "overview": "Mock overview",
            "experiences": [],
        }
        work = make_render_work(
            cv_data, tmp_path / "template.docx", tmp_path / "output.docx"
        )
        result = renderer.render(work)
        assert renderer.last_render["cv_data"] == cv_data
        assert result.output.exists()
        assert result.output.read_text() == "Mock rendered content"


class TestRendererWithExternalParameters:
    """Tests for passing template and data as external parameters."""

    def test_renderer_accepts_external_cv_data(self, tmp_path, make_render_work):
        """Renderer can accept CV data from external sources."""

        class MockCVRenderer(CVRenderer):
            def render(self, work):
                with work.input.open("r", encoding="utf-8") as f:
                    cv_data = json.load(f)
                assert "identity" in cv_data
                assert "sidebar" in cv_data
                work.output.parent.mkdir(parents=True, exist_ok=True)
                work.output.write_text(json.dumps(cv_data))
                return work

        # Simulate external data source
        external_data = {
            "identity": {
                "title": "Senior Developer",
                "full_name": "Jane Doe",
                "first_name": "Jane",
                "last_name": "Doe",
            },
            "sidebar": {
                "languages": ["Python", "JavaScript"],
                "tools": ["Docker"],
                "industries": ["Tech"],
                "spoken_languages": ["English"],
                "academic_background": ["BS CS"],
            },
            "overview": "Experienced developer",
            "experiences": [],
        }

        renderer = MockCVRenderer()
        work = make_render_work(
            external_data, tmp_path / "template.docx", tmp_path / "output.docx"
        )
        result = renderer.render(work)

        # Verify output contains the external data
        rendered_data = json.loads(result.output.read_text())
        assert rendered_data["identity"]["full_name"] == "Jane Doe"
        assert "Python" in rendered_data["sidebar"]["languages"]

    def test_renderer_accepts_external_template_path(self, tmp_path, make_render_work):
        """Renderer can accept template path from external source."""

        class MockCVRenderer(CVRenderer):
            def render(self, work):
                template_path = work.config.render.template
                assert template_path.name == "custom_template.docx"
                work.output.parent.mkdir(parents=True, exist_ok=True)
                work.output.write_text(f"Used template: {template_path}")
                return work

        cv_data = {
            "identity": {
                "title": "",
                "full_name": "",
                "first_name": "",
                "last_name": "",
            },
            "sidebar": {},
            "overview": "",
            "experiences": [],
        }

        # Simulate external template path
        custom_template = tmp_path / "templates" / "custom_template.docx"

        renderer = MockCVRenderer()
        work = make_render_work(cv_data, custom_template, tmp_path / "output.docx")
        result = renderer.render(work)

        assert "custom_template.docx" in result.output.read_text()


class TestDocxCVRendererFullRendering:
    """Tests for complete rendering flow with DocxCVRenderer."""

    def test_render_with_valid_template_creates_output(
        self, tmp_path, monkeypatch, make_render_work
    ):
        """Test successful rendering with a valid template."""
        from cvextract.renderers import docx_renderer

        # Mock DocxTemplate
        class MockTemplate:
            def __init__(self, path):
                self.path = path
                self.rendered_data = None
                self.rendered_autoescape = None
                self.saved_path = None

            def render(self, data, autoescape=False):
                self.rendered_data = data
                self.rendered_autoescape = autoescape

            def save(self, path):
                self.saved_path = path

        mock_tpl = MockTemplate("")
        monkeypatch.setattr(docx_renderer, "DocxTemplate", lambda path: mock_tpl)

        # Create a valid template file
        template_path = tmp_path / "template.docx"
        template_path.write_text("dummy template")

        output_path = tmp_path / "output" / "result.docx"

        cv_data = {
            "identity": {
                "title": "Engineer",
                "full_name": "Test User",
                "first_name": "Test",
                "last_name": "User",
            },
            "sidebar": {
                "languages": ["Python"],
                "tools": ["Docker"],
                "industries": ["Tech"],
                "spoken_languages": ["English"],
                "academic_background": ["BS CS"],
            },
            "overview": "Test overview",
            "experiences": [
                {
                    "heading": "2020-Present | Engineer",
                    "description": "Working on stuff",
                    "bullets": ["Built things", "Fixed bugs"],
                    "environment": ["Python", "AWS"],
                }
            ],
        }

        renderer = DocxCVRenderer()
        work = make_render_work(cv_data, template_path, output_path)
        result = renderer.render(work)

        # Verify output
        assert result.output == output_path
        assert result.output.parent.exists()
        assert mock_tpl.rendered_data is not None
        assert mock_tpl.rendered_autoescape is True
        assert mock_tpl.saved_path == str(output_path)

    def test_render_sanitizes_data(self, tmp_path, monkeypatch, make_render_work):
        """Test that render sanitizes data before rendering."""
        from cvextract.renderers import docx_renderer

        # Mock DocxTemplate
        class MockTemplate:
            def __init__(self, path):
                self.rendered_data = None

            def render(self, data, autoescape=False):
                self.rendered_data = data

            def save(self, path):
                pass

        mock_tpl = MockTemplate("")
        monkeypatch.setattr(docx_renderer, "DocxTemplate", lambda path: mock_tpl)

        template_path = tmp_path / "template.docx"
        template_path.write_text("dummy")
        output_path = tmp_path / "output.docx"

        # Data with characters that need sanitization
        cv_data = {
            "identity": {
                "title": "Engineer\u00a0",  # Non-breaking space
                "full_name": "Test\u00adUser",  # Soft hyphen
                "first_name": "Test",
                "last_name": "User",
            },
            "sidebar": {},
            "overview": "Overview with special chars\u00a0\u00ad",
            "experiences": [],
        }

        renderer = DocxCVRenderer()
        work = make_render_work(cv_data, template_path, output_path)
        renderer.render(work)

        # Verify sanitization was called (data should be different after sanitization)
        assert mock_tpl.rendered_data is not None
        # The sanitized data should have replaced non-breaking spaces and soft hyphens
        assert "\u00a0" not in str(mock_tpl.rendered_data)

    def test_render_with_directory_template_raises_value_error(
        self, tmp_path, make_render_work
    ):
        """Test that render raises ValueError when template is a directory."""
        renderer = DocxCVRenderer()

        # Create a directory instead of a file
        template_dir = tmp_path / "template.docx"
        template_dir.mkdir()

        cv_data = {
            "identity": {
                "title": "",
                "full_name": "",
                "first_name": "",
                "last_name": "",
            },
            "sidebar": {},
            "overview": "",
            "experiences": [],
        }

        work = make_render_work(cv_data, template_dir, tmp_path / "output.docx")
        with pytest.raises(ValueError, match="not a file"):
            renderer.render(work)

    def test_render_creates_output_directory_if_needed(
        self, tmp_path, monkeypatch, make_render_work
    ):
        """Test that render creates output directory if it doesn't exist."""
        from cvextract.renderers import docx_renderer

        # Mock DocxTemplate
        class MockTemplate:
            def render(self, data, autoescape=False):
                pass

            def save(self, path):
                pass

        monkeypatch.setattr(docx_renderer, "DocxTemplate", lambda path: MockTemplate())

        template_path = tmp_path / "template.docx"
        template_path.write_text("dummy")

        # Output path in a non-existent directory
        output_path = tmp_path / "nested" / "deep" / "output.docx"
        assert not output_path.parent.exists()

        cv_data = {
            "identity": {
                "title": "",
                "full_name": "",
                "first_name": "",
                "last_name": "",
            },
            "sidebar": {},
            "overview": "",
            "experiences": [],
        }

        renderer = DocxCVRenderer()
        work = make_render_work(cv_data, template_path, output_path)
        result = renderer.render(work)

        assert result.output == output_path
        assert result.output.parent.exists()


class TestDocxCVRendererIntegration:
    """Integration tests with real DOCX files (minimal mocking)."""

    def test_render_with_real_docx_template(
        self, real_docx_template, tmp_path, make_render_work
    ):
        """Test rendering with an actual DOCX template file."""
        output_path = tmp_path / "output.docx"

        cv_data = {
            "identity": {
                "title": "Senior Software Engineer",
                "full_name": "John Doe",
                "first_name": "John",
                "last_name": "Doe",
            },
            "sidebar": {
                "languages": ["Python", "JavaScript"],
                "tools": ["Docker", "Kubernetes"],
                "industries": ["Finance", "Technology"],
                "spoken_languages": ["English", "Spanish"],
                "academic_background": ["BS Computer Science"],
            },
            "overview": "Experienced software engineer with 10+ years in the industry.",
            "experiences": [
                {
                    "heading": "2020-Present | Senior Engineer at Tech Corp",
                    "description": "Leading development of cloud infrastructure",
                    "bullets": [
                        "Architected microservices",
                        "Improved performance by 40%",
                    ],
                    "environment": ["Python", "AWS", "Docker"],
                }
            ],
        }

        renderer = DocxCVRenderer()
        work = make_render_work(cv_data, real_docx_template, output_path)
        result = renderer.render(work)

        # Verify output file was created
        assert result.output == output_path
        assert result.output.exists()
        assert result.output.stat().st_size > 0

        # Verify it's a valid DOCX file by trying to open it
        from docx import Document

        doc = Document(str(result.output))
        text = "\n".join([p.text for p in doc.paragraphs])

        # Verify the data was rendered into the template
        assert "John Doe" in text
        assert "Senior Software Engineer" in text
        assert "Experienced software engineer" in text
